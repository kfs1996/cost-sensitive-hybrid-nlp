import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    
def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

doc = Document()

# Title
title = doc.add_heading('Chapter Overview: Phase 2 - Cost-Aware Ensemble Stacking', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_paragraph(doc, "This document provides a highly detailed, deeply technical breakdown of Phase 2 of the research methodology. While Phase 1 demonstrated the efficacy of Feature Fusion (concatenating sparse lexical and dense semantic features), Phase 2 investigates Ensemble Stacking. The primary objective of Phase 2 is to determine if a \"Meta-Classifier\" can learn to intelligently resolve disagreements between isolated Base Experts (TF-IDF and Sentence-BERT), particularly when prioritizing rare Non-Functional Requirements (NFRs) via cost-aware weighting.")

add_paragraph(doc, "Phase 2 is divided into two distinct architectural approaches: Phase 2.1 (Hard Stacking) and Phase 2.2 (Soft Stacking).")

# Section 1
add_heading(doc, '1. The Core Architecture: Nested Cross-Validation', 1)
add_paragraph(doc, "Before discussing the differences between Hard and Soft stacking, it is critical to understand how the stacking was trained. If a Meta-Classifier is trained on the same data that the Base Experts were trained on, it will suffer from massive data leakage (overfitting). To prevent this, we engineered a rigorous Nested 5-Fold Cross-Validation loop.")

add_heading(doc, 'How it was engineered:', 2)
add_bullet(doc, "Outer Loop (Evaluation): The dataset is split into 5 folds. 4 folds are used for training the entire ensemble, and 1 fold is held completely unseen for the final test.")
add_bullet(doc, "Inner Loop (Meta-Training): Inside the 4 training folds, the data is split again. The Base Experts (TF-IDF and SBERT) are trained on a subset, and they make predictions on the remaining subset.")
add_bullet(doc, "Out-of-Fold Predictions: Because the Base Experts had never seen that remaining subset, their predictions are \"honest.\" These honest predictions become the actual training data (the features) for the Meta-Classifier.")

add_paragraph(doc, "(Note: See Markdown artifact for the Mermaid Architecture Diagram for Nested CV)", italic=True)

# Section 2
add_heading(doc, '2. Phase 2.1: Hard Stacking (The Information Bottleneck)', 1)
add_heading(doc, 'What We Did', 2)
add_paragraph(doc, "In Hard Stacking, the Base Experts are forced to make a definitive, discrete decision before passing their output to the Meta-Classifier. We used the StackingClassifier from scikit-learn, but deliberately stripped away the probability outputs.")

add_heading(doc, 'How We Did It (Technical Implementation)', 2)
add_bullet(doc, "Expert Processing: The raw text was passed to TF-IDF and SBERT independently. Both experts evaluated the text against their trained Support Vector Machines (SVMs).")
add_bullet(doc, "Hard Voting (predict): We forced the SVMs to use the standard predict() function, which outputs a single integer representing the winning class (e.g., 1 for Security, 2 for Usability).")
add_bullet(doc, "Meta-Classification: The Meta-Classifier (Logistic Regression, Random Forest, or SVM) received an array of discrete integers [1, 3] as its only input features. It attempted to learn which expert to trust based on historical accuracy.")

add_heading(doc, 'The Results & The Bottleneck Discovery', 2)
add_paragraph(doc, "Phase 2.1 yielded unexpectedly poor results across both datasets. Even after applying massive Grid Search hyperparameter tuning, Phase 2.1 failed to cross the 90% threshold on FNFC (peaking at 89.24% with Random Forest) and performed abysmally on minority-class detection (Macro-F1 of ~10% for SVM and LogReg).")

p = doc.add_paragraph()
run1 = p.add_run("The Discovery: ")
run1.bold = True
run2 = p.add_run("We identified a severe Information Bottleneck. By forcing the Base Experts to output rigid integer classifications (1 or 3), all mathematical nuance and confidence metrics were destroyed. If SBERT was 99% confident, and TF-IDF was 51% confident, the Meta-Classifier could not see the difference. It only saw [1, 3]. The Meta-Classifier had no variance to learn from, rendering the stacking architecture fundamentally ineffective for rare-class NFR detection.")

add_paragraph(doc, "(Note: See Markdown artifact for the Mermaid Architecture Diagram for Hard Stacking)", italic=True)

# Section 3
add_heading(doc, '3. Phase 2.2: Soft Stacking (The Breakthrough)', 1)
add_heading(doc, 'What We Did', 2)
add_paragraph(doc, "To resolve the Information Bottleneck, we transitioned to Soft Stacking. Instead of passing discrete votes, the Base Experts pass their continuous probability distributions (or decision function distances).")

add_heading(doc, 'How We Did It (Technical Implementation)', 2)
add_bullet(doc, "Expert Processing: The requirement is passed to TF-IDF and SBERT.")
add_bullet(doc, "Soft Confidence Output (predict_proba or decision_function): Instead of outputting an integer, we instructed the Base Experts to output a continuous array. For example, if there are two classes, TF-IDF outputs [0.51, 0.49] and SBERT outputs [0.01, 0.99].")
add_bullet(doc, "Cost-Aware Meta-Classification: The Meta-Classifier receives the concatenated array of continuous decimal probabilities: [0.51, 0.49, 0.01, 0.99].")
add_bullet(doc, "Grid Search Tuning: We applied GridSearchCV to the Meta-Classifier. We mathematically enforced a class_weight='balanced' (our alpha=0.5 equivalent), forcing the Meta-Classifier to hypersensitize itself to the rarest classes. The Grid Search systematically tested hundreds of internal parameters (like the C penalty parameter in SVM) to find the absolute mathematically optimal configuration for decoding those probability arrays.")

add_heading(doc, 'What We Achieved', 2)
add_paragraph(doc, "Phase 2.2 successfully recovered the lost information and yielded breakthrough results, particularly on the highly complex, multi-class PROMISE dataset.")

add_paragraph(doc, "Dataset 1: FNFC (Binary)", bold=True)
add_bullet(doc, "The Tuned SVM Meta-Classifier achieved 91.24% Accuracy, nearly matching the Phase 1 Feature Fusion ceiling (91.47%).")
add_bullet(doc, "More importantly, it achieved the highest Macro-F1 score (50.22%) of the entire experiment, proving that Soft Stacking is vastly superior at identifying the rare minority classes in binary datasets.")

add_paragraph(doc, "Dataset 2: PROMISE (Multi-Class)", bold=True)
add_bullet(doc, "The complexity of Soft Stacking proved to be perfectly suited for the chaotic, 11-class PROMISE dataset.")
add_bullet(doc, "Phase 1 Feature Fusion peaked at 80.18%.")
add_bullet(doc, "The Tuned Phase 2.2 SVM Meta-Classifier broke the ceiling, achieving an unprecedented 81.52% Accuracy.")
add_bullet(doc, "This proves conclusively that for complex, multi-class text environments, isolated experts feeding continuous probability distributions into a tuned, cost-aware Meta-Classifier yields the highest mathematical performance.")

add_paragraph(doc, "(Note: See Markdown artifact for the Mermaid Architecture Diagram for Soft Stacking)", italic=True)

# Section 4
add_heading(doc, '4. Conclusion & Future Work Context', 1)
add_heading(doc, 'The Final Verdict', 2)
add_bullet(doc, "For simple, binary classification (FNFC), the raw mathematical combination of Feature Fusion (Phase 1) is the most efficient and accurate architecture.")
add_bullet(doc, "For complex, multi-class classification (PROMISE), Cost-Aware Soft Stacking (Phase 2.2) is the undisputed champion.")

add_heading(doc, 'Future Directions (Domain-Specific Embeddings)', 2)
add_paragraph(doc, "While this study utilized a highly efficient general-purpose semantic embedding model (Sentence-BERT: all-MiniLM-L6), future research should evaluate this exact Soft Stacking architecture using domain-specific models such as seBERT (Software Engineering BERT) or RE-BERT (Requirements Engineering BERT). Due to the significant computational overhead of these 110M+ parameter models, investigating whether domain-specific embeddings further elevate the Meta-Classifier's ability to detect rare-class requirements remains an open and highly promising area for future large-scale empirical studies.")

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Phase2_Thesis_Explainer.docx")
doc.save(output_path)
print(f"Successfully generated {output_path}")
